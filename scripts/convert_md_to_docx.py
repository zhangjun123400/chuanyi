"""Convert PRD markdown to a well-formatted Word document.

Usage:
    python scripts/convert_md_to_docx.py <input.md> [output.docx]

If output is omitted, it replaces .md with .docx in the input path.
"""

import re
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

if len(sys.argv) < 2:
    print("Usage: python convert_md_to_docx.py <input.md> [output.docx]")
    sys.exit(1)

INPUT = sys.argv[1]
OUTPUT = sys.argv[2] if len(sys.argv) > 2 else INPUT.rsplit('.md', 1)[0] + '.docx'

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.35
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for level in range(1, 5):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Microsoft YaHei'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    h.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    if level == 1:
        h.font.size = Pt(22)
    elif level == 2:
        h.font.size = Pt(16)
    elif level == 3:
        h.font.size = Pt(13)
    else:
        h.font.size = Pt(11.5)


def add_code_block(text, language=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'SF Mono'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    shading = p.paragraph_format.element.get_or_add_pPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F1F5F9',
        qn('w:val'): 'clear',
    })
    shading.append(shd)
    return p


def add_table_from_md(rows):
    if not rows:
        return
    filtered = []
    for r in rows:
        if all(re.match(r'^[-: ]+$', cell.strip()) for cell in r):
            continue
        filtered.append(r)
    if len(filtered) < 2:
        return
    table = doc.add_table(rows=len(filtered), cols=len(filtered[0]))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row in enumerate(filtered):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = cell_text.strip()
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(1)
                paragraph.paragraph_format.space_before = Pt(1)
                for run in paragraph.runs:
                    run.font.size = Pt(9) if i > 0 else Pt(9.5)
                    run.font.name = 'Microsoft YaHei'
                    if i == 0:
                        run.font.bold = True
    doc.add_paragraph()


def parse_inline(text):
    parts = []
    remaining = text
    while remaining:
        m = re.search(r'\*\*(.+?)\*\*', remaining)
        if not m:
            parts.append((remaining, False))
            break
        if m.start() > 0:
            parts.append((remaining[:m.start()], False))
        parts.append((m.group(1), True))
        remaining = remaining[m.end():]
    return parts


def add_paragraph_with_bold(text):
    p = doc.add_paragraph()
    for txt, is_bold in parse_inline(text):
        run = p.add_run(txt)
        run.font.name = 'Microsoft YaHei'
        if is_bold:
            run.font.bold = True
    return p


def add_list_item(text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    for txt, is_bold in parse_inline(text):
        run = p.add_run(txt)
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(10.5)
        if is_bold:
            run.font.bold = True
    return p


# -- Parse markdown --
with open(INPUT, 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
table_buffer = []
in_code_block = False
code_buf = []
code_lang = ""

while i < len(lines):
    line = lines[i].rstrip()

    if line.startswith('```'):
        if in_code_block:
            add_code_block('\n'.join(code_buf), code_lang)
            code_buf = []
            in_code_block = False
        else:
            in_code_block = True
            code_lang = line[3:].strip()
        i += 1
        continue

    if in_code_block:
        code_buf.append(line)
        i += 1
        continue

    if '|' in line and line.strip().startswith('|'):
        cells = [c.strip() for c in line.strip().split('|')[1:-1]]
        table_buffer.append(cells)
        if i + 1 < len(lines) and '|' in lines[i+1] and lines[i+1].strip().startswith('|'):
            i += 1
            continue
        else:
            add_table_from_md(table_buffer)
            table_buffer = []
            i += 1
            continue

    if line.strip() in ('---', '***', '___'):
        doc.add_paragraph('─' * 60)
        i += 1
        continue

    if line.startswith('# ') and not line.startswith('## '):
        doc.add_heading(line[2:].strip(), level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
    elif line.startswith('#### '):
        doc.add_heading(line[5:].strip(), level=4)
    elif re.match(r'^- |^  - ', line):
        text = re.sub(r'^  - |^- ', '', line)
        add_list_item(text)
    elif not line.strip():
        doc.add_paragraph()
    else:
        add_paragraph_with_bold(line)

    i += 1

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
